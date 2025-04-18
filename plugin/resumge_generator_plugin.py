from flask import Flask, request, jsonify
from io import BytesIO
import base64
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from supabase import create_client, Client
import uuid
from datetime import datetime
import json

# Supabase Configuration
SUPABASE_URL = ""
SUPABASE_KEY = ""
BUCKET_NAME = ""
SERVICE_ROLE = ""

# Initialize Supabase Client
supabase: Client = create_client(SUPABASE_URL, SERVICE_ROLE)

app = Flask(__name__)

def create_ats_resume(user_data):
    doc = Document()

    # Set Narrow Margins (0.5 inches for all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    def add_section_title(title):
        """Add a section title with a horizontal line below."""
        if not title:
            return

        p = doc.add_paragraph(title)
        p.runs[0].bold = True
        p.runs[0].font.name = "Times New Roman"
        p.runs[0].font.size = Pt(10)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        border = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        border.append(bottom)
        p._element.get_or_add_pPr().append(border)

    def add_entry_with_right_aligned_date(left_text, right_text):
        """Adds a paragraph with left-aligned text and right-aligned text using a tab stop."""
        if not left_text and not right_text:
            return

        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        if left_text:
            run = p.add_run(left_text)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

        p.add_run("\t")

        if right_text:
            run = p.add_run(right_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    def add_hyperlink(paragraph, text, url):
        """Add a clickable hyperlink to a paragraph."""
        if not text or not url:
            return

        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = paragraph.add_run(text)
        new_run.font.size = Pt(10)
        new_run.font.name = 'Times New Roman'
        new_run.font.color.rgb = RGBColor(0, 0, 255)
        new_run.font.underline = True

        r = new_run._element
        hyperlink.append(r)
        paragraph._element.append(hyperlink)

    # Header - Name
    if user_data.get("name"):
        name = user_data["name"].upper()
    else:
        name = "Your Name"
    p = doc.add_paragraph(name)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    # Header - Contact Info with Hyperlinks
    contact_info = []
    if "phone" in user_data:
        contact_info.append(user_data["phone"])
    if "email" in user_data:
        contact_info.append(user_data["email"])

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if contact_info:
        p.add_run("  |  ".join(contact_info) + "  |  ")

    # Add Hyperlinks
    for text, url in [
        ("LinkedIn", user_data.get("linked_in")),
        ("GitHub", user_data.get("github")),
        ("Portfolio", user_data.get("portfolio_website"))
    ]:
        if url:
            add_hyperlink(p, text, f"https://{url}")
            p.add_run("  |  ")

    # Section: INTRODUCTION
    intro = user_data.get("introduction_paragraph")
    if intro:
        add_section_title("INTRODUCTION")
        p = doc.add_paragraph(intro)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.runs[0].font.name = "Times New Roman"
        p.runs[0].font.size = Pt(10)

    # Section: EDUCATION
    education_list = user_data.get("education", [])
    if education_list and education_list != [{}]:
        if [edu for edu in education_list[0].values() if edu]:
            add_section_title("EDUCATION")
            for edu in education_list:
                add_entry_with_right_aligned_date(edu.get("institution"), edu.get("period"))
                if "course_name" in edu:
                    p = doc.add_paragraph(edu["course_name"])
                    p.runs[0].italic = True
                    p.runs[0].font.name = "Times New Roman"
                    p.runs[0].font.size = Pt(10)
                    # Set space_after on the paragraph_format
                    p.paragraph_format.space_after = Pt(0)  # This will remove the space after the paragraph


                for field in ["results", "related_coursework"]:
                    if field in edu:
                        doc.add_paragraph(edu[field], style="List Bullet").runs[0].font.size = Pt(10)

    # Section: EXPERIENCE
    experience_list = user_data.get("experience", [])
    if experience_list and experience_list != [{}]:
        if [edu for edu in experience_list[0].values() if edu]:
            add_section_title("EXPERIENCE")
            for exp in experience_list:
                add_entry_with_right_aligned_date(exp.get("company"), exp.get("period"))
                for field in ["job_description", "related_knowledge"]:
                    if field in exp:
                        doc.add_paragraph(exp[field], style="List Bullet").runs[0].font.size = Pt(10)

    # Section: PROJECTS
    projects_list = user_data.get("projects", [])
    if projects_list and projects_list != [{}]:
        if [edu for edu in projects_list[0].values() if edu]:
            add_section_title("PROJECTS")
            for proj in projects_list:
                add_entry_with_right_aligned_date(proj.get("project_name"), proj.get("period"))
                for field in ["project_description", "related_knowledge"]:
                    if field in proj:
                        doc.add_paragraph(proj[field], style="List Bullet").runs[0].font.size = Pt(10)

    # Section: ACHIEVEMENTS
    achievements_list = user_data.get("achievements", [])
    if achievements_list and achievements_list != [{}]:
        if [edu for edu in achievements_list[0].values() if edu]:
            add_section_title("ACHIEVEMENTS")
            for ach in achievements_list:
                add_entry_with_right_aligned_date(ach.get("ach_title"), ach.get("ach_date"))
                if "ach_description" in ach:
                    doc.add_paragraph(ach["ach_description"], style="List Bullet").runs[0].font.size = Pt(10)

    # Section: SKILLS
    skills_list = user_data.get("skills", [])
    if skills_list and skills_list != [{}]:
        if [edu for edu in skills_list[0].values() if edu]:
            add_section_title("SKILLS")

            for category in skills_list:
                skill_name = category.get("skill_name", "")  # Get skill category
                skills = category.get("skills", "")  # Get skills string

                if skill_name and skills:
                    split_skills = skills.split(", ")  # Use " | " as separator
                    skills_text = f"{skill_name}: {' | '.join(split_skills)}"
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(skills_text).font.name = "Times New Roman"
                    p.runs[0].font.size = Pt(10)

    # Save document to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def upload_to_supabase(file_data, file_name):
    """Uploads the DOCX file to Supabase Storage and returns the public URL."""
    response = supabase.storage.from_(BUCKET_NAME).upload(file_name, file_data.getvalue(), {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"})

    # Successful upload
    public_url = f"{SUPABASE_URL}/storage/v1/object/public/{BUCKET_NAME}/{file_name}"
    return public_url

def generate_unique_filename(file_name):
    unique_id = uuid.uuid4().hex  # Or use datetime.now().timestamp() for a timestamp-based name
    base_name, ext = file_name.rsplit('.', 1)
    return f"{base_name}_{unique_id}.{ext}"

@app.route('/generate_resume', methods=['POST'])
def generate_resume():
    user_data = request.get_json()
    
    if isinstance(user_data, str):  # Handle string-encoded JSON
        user_data = json.loads(user_data)
        
        print("bye")
        
    doc_file = create_ats_resume(user_data)

    # Define file name
    file_name = generate_unique_filename("resume.docx")

    # Upload to Supabase
    file_url = upload_to_supabase(doc_file, file_name)

    if file_url:
        return jsonify({"message": "Resume uploaded successfully", "download_url": file_url})
    else:
        return jsonify({"error": "Failed to upload"}), 500

if __name__ == '__main__':
    app.run(debug=True)

# {
#   "name": "Yoo Jae Yi",
#   "phone": "+1 123 456 7890",
#   "email": "jay@example.com",
#   "linked_in": "linkedin.com/in/jay",
#   "github": "github.com/jay",
#   "portfolio_website": "jay.com",
#   "introduction_paragraph": "An enthusiastic software engineer with experience in developing innovative applications. Passionate about coding and solving real-world problems with technology.",
#   "education": [
#     {
#       "institution": "Harvard University",
#       "period": "2015 - 2019",
#       "course_name": "Bachelor of Computer Science",
#       "results": "3.9 GPA",
#       "related_coursework": "Data Structures, Algorithms, Machine Learning"
#     },
#     {
#       "institution": "Stanford University",
#       "period": "2020 - 2022",
#       "course_name": "Master of Software Engineering",
#       "results": "4.0 GPA",
#       "related_coursework": "Distributed Systems, Artificial Intelligence"
#     }
#   ],
#   "experience": [
#     {
#       "company": "Tech Corp",
#       "period": "2021 - Present",
#       "job_description": "Developed web applications using React and Node.js. Led the implementation of RESTful APIs.",
#       "related_knowledge": "React, Node.js, RESTful APIs, MongoDB"
#     },
#     {
#       "company": "Dev Solutions",
#       "period": "2019 - 2021",
#       "job_description": "Collaborated in a team to build a full-stack e-commerce platform. Worked on both frontend and backend.",
#       "related_knowledge": "HTML, CSS, JavaScript, Python, Django"
#     }
#   ],
#   "projects": [
#     {
#       "project_name": "Personal Portfolio Website",
#       "period": "2021",
#       "project_description": "Built a personal portfolio website to showcase various projects and skills. Integrated with GitHub for project management.",
#       "related_knowledge": "HTML, CSS, JavaScript, GitHub"
#     },
#     {
#       "project_name": "Weather App",
#       "period": "2020",
#       "project_description": "Developed a weather forecasting app using React. Integrated OpenWeather API for real-time data.",
#       "related_knowledge": "React, API Integration"
#     }
#   ],
#   "achievements": [
#     {
#       "achievement": "Best Developer Award",
#       "date_or_period": "2022",
#       "achievement_description": "Received the Best Developer Award for outstanding performance at Tech Corp."
#     },
#     {
#       "achievement": "Hackathon Winner",
#       "date_or_period": "2021",
#       "achievement_description": "Won the first prize in a company-wide hackathon for creating an innovative app idea."
#     }
#   ],  
#   "skills": [
#     {
#         "skill_name": "Programming Languages",
#         "skills": "Python, JavaScript, Java, C++"
#     },
#     {
#         "skill_name": "Frameworks",
#         "skills": "React, Node.js, Django, Flask"
#     }
#   ]
# }

# {
#   "name": "[Full Name]",
#   "phone": "[Phone Number]",
#   "email": "[Email Address]",
#   "linked_in": "[LinkedIn Profile URL]",
#   "github": "[GitHub Profile URL]",
#   "portfolio_website": "[Portfolio Website URL]",
#   "introduction_paragraph": "[Brief personal summary highlighting experience, skills, and passion for the field]",
  
#   "education": [
#     {
#       "institution": "[University or Institution Name]",
#       "period": "[Start Year - End Year]",
#       "course_name": "[Degree or Course Name]",
#       "results": "[GPA or Notable Academic Achievement]",
#       "related_coursework": "[Key subjects or coursework relevant to the field]"
#     },
#     {
#       "institution": "[University or Institution Name]",
#       "period": "[Start Year - End Year]",
#       "course_name": "[Degree or Course Name]",
#       "results": "[GPA or Notable Academic Achievement]",
#       "related_coursework": "[Key subjects or coursework relevant to the field]"
#     }
#   ],
  
#   "experience": [
#     {
#       "company": "[Company Name]",
#       "period": "[Start Year - End Year or Present]",
#       "job_description": "[Key responsibilities and achievements in the role]",
#       "related_knowledge": "[Relevant technologies, frameworks, or methodologies used]"
#     },
#     {
#       "company": "[Company Name]",
#       "period": "[Start Year - End Year or Present]",
#       "job_description": "[Key responsibilities and achievements in the role]",
#       "related_knowledge": "[Relevant technologies, frameworks, or methodologies used]"
#     }
#   ],
  
#   "projects": [
#     {
#       "project_name": "[Project Title]",
#       "period": "[Year or Duration]",
#       "project_description": "[Brief summary of the project, its goals, and the role played]",
#       "related_knowledge": "[Technologies, frameworks, or tools used]"
#     },
#     {
#       "project_name": "[Project Title]",
#       "period": "[Year or Duration]",
#       "project_description": "[Brief summary of the project, its goals, and the role played]",
#       "related_knowledge": "[Technologies, frameworks, or tools used]"
#     }
#   ],
  
#   "achievements": [
#     {
#       "achievement": "[Award or Recognition Title]",
#       "date_or_period": "[Year or Date]",
#       "achievement_description": "[Brief explanation of the achievement and its significance]"
#     },
#     {
#       "achievement": "[Award or Recognition Title]",
#       "date_or_period": "[Year or Date]",
#       "achievement_description": "[Brief explanation of the achievement and its significance]"
#     }
#   ],
  
#   "skills": [
#     {
#       "skill_name": "[skill name]",
#       "skills": "[skill 1], [skill 2], [skill 3]"
#     }
#   ]
# }
