from flask import Flask, request, jsonify
from io import BytesIO
import base64
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from supabase import create_client, Client
import uuid
from datetime import datetime
import json

# Supabase Configuration
SUPABASE_URL = "https://zasleszppwndgrajepva.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6Inphc2xlc3pwcHduZGdyYWplcHZhIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NDE3NTcxNTksImV4cCI6MjA1NzMzMzE1OX0.01lXJaFUM55DK4EWslmdosboNhGRETo2fbsFtnuRMck"
BUCKET_NAME = "test-bucket"
SERVICE_ROLE = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6Inphc2xlc3pwcHduZGdyYWplcHZhIiwicm9sZSI6InNlcnZpY2Vfcm9sZSIsImlhdCI6MTc0MTc1NzE1OSwiZXhwIjoyMDU3MzMzMTU5fQ.yFjXTwQDFqFj97YquJu8dMQUrlTOJeMG09FiVpz4EH0"

# Initialize Supabase Client
supabase: Client = create_client(SUPABASE_URL, SERVICE_ROLE)

app = Flask(__name__)

def create_ats_resume(user_data):
    doc = Document()

    # Set Narrow Margins (0.5 inches for all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    def add_section_title(title):
        """Add a section title with a horizontal line below."""
        if not title:
            return

        p = doc.add_paragraph(title)
        p.runs[0].bold = True
        p.runs[0].font.name = "Times New Roman"
        p.runs[0].font.size = Pt(10)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        border = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        border.append(bottom)
        p._element.get_or_add_pPr().append(border)

    def add_entry_with_right_aligned_date(left_text, right_text):
        """Adds a paragraph with left-aligned text and right-aligned text using a tab stop."""
        if not left_text and not right_text:
            return

        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        if left_text:
            run = p.add_run(left_text)
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

        p.add_run("\t")

        if right_text:
            run = p.add_run(right_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    def add_hyperlink(paragraph, text, url):
        """Add a clickable hyperlink to a paragraph."""
        if not text or not url:
            return

        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = paragraph.add_run(text)
        new_run.font.size = Pt(10)
        new_run.font.name = 'Times New Roman'
        new_run.font.color.rgb = RGBColor(0, 0, 255)
        new_run.font.underline = True

        r = new_run._element
        hyperlink.append(r)
        paragraph._element.append(hyperlink)

    # Header - Name
    if user_data.get("name"):
        name = user_data["name"].upper()
    else:
        name = "Your Name"
    p = doc.add_paragraph(name)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    # Header - Contact Info with Hyperlinks
    contact_info = []
    if "phone" in user_data:
        contact_info.append(user_data["phone"])
    if "email" in user_data:
        contact_info.append(user_data["email"])

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if contact_info:
        p.add_run("  |  ".join(contact_info) + "  |  ")

    # Add Hyperlinks
    for text, url in [
        ("LinkedIn", user_data.get("linked_in")),
        ("GitHub", user_data.get("github")),
        ("Portfolio", user_data.get("portfolio_website"))
    ]:
        if url:
            add_hyperlink(p, text, f"https://{url}")
            p.add_run("  |  ")

    # Section: INTRODUCTION
    intro = user_data.get("introduction_paragraph")
    if intro:
        add_section_title("INTRODUCTION")
        p = doc.add_paragraph(intro)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.runs[0].font.name = "Times New Roman"
        p.runs[0].font.size = Pt(10)

    # Section: EDUCATION
    education_list = user_data.get("education", [])
    if education_list:
        add_section_title("EDUCATION")
        for edu in education_list:
            add_entry_with_right_aligned_date(edu.get("institution"), edu.get("period"))
            if "course_name" in edu:
                p = doc.add_paragraph(edu["course_name"])
                p.runs[0].italic = True
                p.runs[0].font.name = "Times New Roman"
                p.runs[0].font.size = Pt(10)

            for field in ["results", "related_coursework"]:
                if field in edu:
                    doc.add_paragraph(edu[field], style="List Bullet").runs[0].font.size = Pt(10)

    # Section: EXPERIENCE
    experience_list = user_data.get("experience", [])
    if experience_list:
        add_section_title("EXPERIENCE")
        for exp in experience_list:
            add_entry_with_right_aligned_date(exp.get("company"), exp.get("period"))
            for field in ["job_description", "related_knowledge"]:
                if field in exp:
                    doc.add_paragraph(exp[field], style="List Bullet").runs[0].font.size = Pt(10)

    # Section: PROJECTS
    projects_list = user_data.get("projects", [])
    if projects_list:
        add_section_title("PROJECTS")
        for proj in projects_list:
            add_entry_with_right_aligned_date(proj.get("project_name"), proj.get("period"))
            for field in ["project_description", "related_knowledge"]:
                if field in proj:
                    doc.add_paragraph(proj[field], style="List Bullet").runs[0].font.size = Pt(10)

    # Section: ACHIEVEMENTS
    achievements_list = user_data.get("achievements", [])
    if achievements_list:
        add_section_title("ACHIEVEMENTS")
        for ach in achievements_list:
            add_entry_with_right_aligned_date(ach.get("achievement"), ach.get("date/period"))
            if "achievement_description" in ach:
                doc.add_paragraph(ach["achievement_description"], style="List Bullet").runs[0].font.size = Pt(10)

    # Section: SKILLS
    skills_list = user_data.get("skills", [])
    if skills_list:
        add_section_title("SKILLS")
        for category in skills_list:
            for skill, skills in category.items():
                if skills:
                    skills_text = f"{skill}: {' | '.join(skills)}"
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(skills_text).font.name = "Times New Roman"
                    p.runs[0].font.size = Pt(10)

    # Save document to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def upload_to_supabase(file_data, file_name):
    """Uploads the DOCX file to Supabase Storage and returns the public URL."""
    response = supabase.storage.from_(BUCKET_NAME).upload(file_name, file_data.getvalue(), {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"})

    # Successful upload
    public_url = f"{SUPABASE_URL}/storage/v1/object/public/{BUCKET_NAME}/{file_name}"
    return public_url

def generate_unique_filename(file_name):
    unique_id = uuid.uuid4().hex  # Or use datetime.now().timestamp() for a timestamp-based name
    base_name, ext = file_name.rsplit('.', 1)
    return f"{base_name}_{unique_id}.{ext}"

@app.route('/generate_resume', methods=['POST'])
def generate_resume():
# Output: <class 'str'>
    user_data = request.get_json()
    
    print("hi")
    print(type(user_data))
    
    
    if isinstance(user_data, str):  # Handle string-encoded JSON
        user_data = json.loads(user_data)
        
        print("bye")
        
    doc_file = create_ats_resume(user_data)

    # Define file name
    file_name = generate_unique_filename("resume.docx")

    # Upload to Supabase
    file_url = upload_to_supabase(doc_file, file_name)

    if file_url:
        return jsonify({"message": "Resume uploaded successfully", "download_url": file_url})
    else:
        return jsonify({"error": "Failed to upload"}), 500

if __name__ == '__main__':
    app.run(debug=True)

# dummy_json_data = {
#   "name": "John Doe",
#   "phone": "+1 123 456 7890",
#   "email": "johndoe@example.com",
#   "portfolio_website": "johndoe.com",
#   "introduction_paragraph": "Detail-oriented software engineer with 5 years of experience in full-stack web development. Passionate about building scalable applications and improving system performance.",
  
#   "experience": [
#     {
#       "company": "Tech Solutions Inc.",
#       "period": "2020 - Present",
#       "job_description": "Developed and maintained web applications using React and Node.js.",
#       "related_knowledge": "JavaScript, React, Node.js, Express, MongoDB"
#     },
#     {
#       "company": "Web Innovations Ltd.",
#       "period": "2018 - 2020",
#       "job_description": "Designed and optimized RESTful APIs for enterprise applications.",
#       "related_knowledge": "RESTful APIs, Python, Flask, PostgreSQL"
#     }
#   ],
#   "projects": [
#     {
#       "project_name": "E-commerce Platform",
#       "project_description": "Developed a fully functional e-commerce platform with secure payment integration.",
#       "related_knowledge": "React, Redux, Stripe API, Firebase"
#     },
#     {
#       "project_name": "Chatbot Application",
#       "period": "2020",
#       "project_description": "Built an AI-powered chatbot for customer support automation.",
#       "related_knowledge": "Python, TensorFlow, NLP"
#     }
#   ],
#   "achievements": [
#     {
#       "achievement": "Hackathon Winner",
#       "date/period": "2019",
#       "achievement_description": "Won first place in a national hackathon for developing an AI-driven recommendation system."
#     },
#     {
#       "achievement": "Certified AWS Solutions Architect",
#       "date/period": "2021",
#       "achievement_description": "Earned AWS Solutions Architect certification demonstrating cloud expertise."
#     }
#   ],
#   "skills": [
#       {"frontend development": ["HTML", "CSS", "JavaScript", "React", "Angular"]},
#       {"backend development": ["Node.js", "Express", "Python", "Django", "Flask"]},
#       {"databases": ["MongoDB", "PostgreSQL", "MySQL", "Firebase"]},
#       {"cloud technologies": ["AWS", "Google Cloud", "Azure"]},
#       {"machine learning": ["TensorFlow", "Scikit-learn", "Keras", "PyTorch"]}
#   ]
# }


# # Example Usage:
# create_ats_resume(dummy_json_data)
